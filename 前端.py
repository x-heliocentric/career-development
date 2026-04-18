#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模方职航 - 最终完整版（包含岗位画像页面）
"""

import streamlit as st
import requests
import json
import time
import io
import re
import os
import networkx as nx
from docx import Document
from typing import Dict, List, Any
import plotly.graph_objects as go
import plotly.express as px

# ========== API配置 ==========
API_URL = "https://gwmr32ycd5.coze.site/stream_run"
# ⚠️ 请确保以下API_TOKEN是最新的，如需更新请联系管理员
API_TOKEN = "eyJhbGciOiJSUzI1NiIsImtpZCI6ImEwYzU4NWYwLTRkNTctNGYzZC05NTY2LWNjYWMyYWU2N2Q5MiJ9.eyJpc3MiOiJodHRwczovL2FwaS5jb3plLmNuIiwiYXVkIjpbIlNwRmFFT243NFZsNGJzVXpGbjdUdEVLeUZMM3hqMnE3Il0sImV4cCI6ODIxMDI2Njg3Njc5OSwiaWF0IjoxNzc1NzIyMTQ4LCJzdWIiOiJzcGlmZmU6Ly9hcGkuY296ZS5jbi93b3JrbG9hZF9pZGVudGl0eS9pZDo3NjI1NjUwNDg0OTMzNTU4MzI2Iiwic3JjIjoiaW5ib3VuZF9hdXRoX2FjY2Vzc190b2tlbl9pZDo3NjI2NjY4NTUzMzExMDkyNzc5In0.dZwVqwuVteCk1e7WY44IiTzRKM-8zizoSwJ6rSam2VvEQPV_KkKAs-DrA2j1wx_xiJsLjNlmhcklMVf1GcnzQQfuww7rKOPNuLDHGy9h5RkWggRiSTxMNM6JwtvNdbnEKzrOx9pK4tT3yifdoBqsXVLLx_JN6CT-_iX7wzkxSWxMZeteN6Bbvrj1mdkmJS48ljkhhpx0o4NrrG5Mhsq8DS1xVJ7JZtFiYakOBJTFToQQcRLKgKCwrprxwJD_XVTtMeixSmBNIS2WoradLz1SlWO8Qa-H1qeSGgx0uE3pCMNBW-aqwXubfPpiehITGvtqS-0Wtuzxz8tBTKlDQSLfhQ"
PROJECT_ID = "7625643214413938724"

# ========== 页面配置 ==========
st.set_page_config(
    page_title="模方职航",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 设置应用Logo（Streamlit 1.12+支持）
try:
    st.logo("logo.png")
except Exception:
    # 如果st.logo()不可用或logo.png不存在，忽略错误
    pass

# ========== 初始化状态 ==========
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'session_id' not in st.session_state:
    st.session_state.session_id = f"FHTjpkOwLO_oN0D0gb8Tr_{int(time.time())}"
if 'page' not in st.session_state:
    st.session_state.page = "💬 对话助手"

# ========== 岗位数据加载 ==========
@st.cache_data(ttl=3600)
def load_job_data():
    """加载岗位数据"""
    workspace_path = os.path.dirname(os.path.abspath(__file__))
    
    try:
        # 加载岗位画像库
        with open(os.path.join(workspace_path, "assets", "job_portraits.json"), 'r', encoding='utf-8') as f:
            job_portraits = json.load(f)
        
        # 加载换岗路径
        with open(os.path.join(workspace_path, "assets", "job_relation_graph.json"), 'r', encoding='utf-8') as f:
            job_relation_graph = json.load(f)
        
        # 加载垂直晋升路径
        with open(os.path.join(workspace_path, "assets", "job_vertical_paths.json"), 'r', encoding='utf-8') as f:
            vertical_paths = json.load(f)
        
        return job_portraits, job_relation_graph, vertical_paths
    except FileNotFoundError:
        return {}, {}, {}

def get_all_jobs(job_portraits):
    """获取所有岗位名称"""
    jobs = [job.get("standard_job_name", "") for job in job_portraits]
    return sorted(list(set(jobs)))

# ========== 岗位能力计算函数 ==========
def calculate_job_scores(job_portrait: Dict) -> Dict[str, float]:
    """
    根据岗位画像动态计算各维度能力得分
    """
    scores = {
        "专业技能": 5.0,
        "沟通能力": 5.0,
        "学习能力": 5.0,
        "创新能力": 5.0,
        "协作能力": 5.0,
        "抗压能力": 5.0,
        "解决问题": 5.0,
        "项目管理": 5.0
    }
    
    skills = job_portrait.get("skills", {})
    common_skills = skills.get("common_skills", [])
    advanced_skills = skills.get("advanced_skills", [])
    job_level = job_portrait.get("job_level", "")
    job_description = job_portrait.get("job_description", "")
    responsibilities = job_portrait.get("responsibilities", [])
    education = job_portrait.get("education_required", [])
    salary = job_portrait.get("salary_range", {}).get("avg", 0)
    
    all_text = f"{job_description} {' '.join(responsibilities)} {' '.join(common_skills)} {' '.join(advanced_skills)}"
    all_text = all_text.lower()
    
    # 专业技能分数
    skill_count = len(common_skills) + len(advanced_skills)
    advanced_ratio = len(advanced_skills) / max(1, len(common_skills))
    skill_score = 5.0 + min(3.0, skill_count * 0.25) + advanced_ratio * 2.0
    scores["专业技能"] = min(10.0, skill_score)
    
    keyword_weights = {
        "沟通能力": [("沟通", 2.0), ("协作", 1.5), ("汇报", 1.5), ("交流", 1.5), ("协调", 1.5)],
        "学习能力": [("学习", 1.5), ("研究", 1.5), ("掌握", 1.0), ("分析", 1.0), ("硕士", 2.5), ("博士", 3.0)],
        "创新能力": [("创新", 2.0), ("设计", 1.5), ("开发", 1.5), ("优化", 1.5), ("研发", 1.5)],
        "协作能力": [("团队", 2.0), ("协作", 2.0), ("配合", 1.5), ("跨部门", 2.0)],
        "抗压能力": [("高压", 3.0), ("加班", 2.0), ("紧急", 2.0), ("目标", 1.0)],
        "解决问题": [("解决", 2.0), ("问题", 1.0), ("分析", 1.5), ("处理", 1.5)],
        "项目管理": [("管理", 2.0), ("项目经理", 3.0), ("规划", 1.5), ("统筹", 2.0)]
    }
    
    for dimension, keywords in keyword_weights.items():
        score = 5.0
        for keyword in keywords:
            if isinstance(keyword, tuple):
                kw, weight = keyword
            else:
                kw, weight = keyword, 1.5
            count = all_text.count(kw.lower())
            if count > 0:
                score += min(weight * count, 3.0)
        scores[dimension] = min(10.0, score)
    
    if "博士" in " ".join(education):
        scores["学习能力"] = min(10.0, scores["学习能力"] + 2.0)
    
    level_match = re.search(r'\d+', str(job_level))
    if level_match and int(level_match.group()) >= 8:
        scores["项目管理"] = min(10.0, scores["项目管理"] + 2.0)
    
    if salary > 30000:
        scores["抗压能力"] = min(10.0, scores["抗压能力"] + 1.5)
    
    return scores

# ========== 岗位可视化函数 ==========
def create_job_radar(job_portrait: Dict) -> go.Figure:
    """创建岗位能力雷达图"""
    scores = calculate_job_scores(job_portrait)
    dimensions = list(scores.keys())
    values = list(scores.values())
    
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=values,
        theta=dimensions,
        fill='toself',
        name='能力评分',
        line_color='rgb(79, 70, 229)'
    ))
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 10])),
        showlegend=True,
        title="岗位能力维度评分（基于岗位要求计算）",
        height=400
    )
    return fig

def create_salary_chart(job_portrait: Dict) -> go.Figure:
    """创建薪资分布图"""
    salary_range = job_portrait.get("salary_range", {})
    min_salary = salary_range.get("min", 0)
    max_salary = salary_range.get("max", 0)
    avg_salary = salary_range.get("avg", 0)
    
    fig = go.Figure(go.Indicator(
        mode = "number+gauge+delta",
        value = avg_salary,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': f"平均月薪 (元)", 'font': {'size': 16}},
        delta = {'reference': 15000, 'increasing': {'color': "RebeccaPurple"}},
        gauge = {
            'axis': {'range': [None, 50000], 'tickwidth': 1, 'tickcolor': "darkblue"},
            'bar': {'color': "darkblue"},
            'bgcolor': "white",
            'borderwidth': 2,
            'bordercolor': "gray",
            'steps': [
                {'range': [0, 10000], 'color': 'lightgray'},
                {'range': [10000, 20000], 'color': 'lightyellow'},
                {'range': [20000, 30000], 'color': 'lightgreen'},
                {'range': [30000, 40000], 'color': 'lightblue'},
                {'range': [40000, 50000], 'color': 'lightpink'}
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': max_salary * 0.8
            }
        }
    ))
    fig.update_layout(
        title=f"薪资范围: {min_salary}-{max_salary} 元/月",
        height=300
    )
    return fig

def create_vertical_path_graph(job_name: str, vertical_paths: Dict) -> go.Figure:
    """创建垂直晋升路径图"""
    jobs_data = vertical_paths.get("jobs", {})
    current_job = jobs_data.get(job_name, {})
    career_path = current_job.get("career_path", {})
    vertical_path = career_path.get("vertical_path", [])
    
    nodes = [{"name": job_name, "level": 0, "type": "current"}]
    edges = []
    
    for i, path in enumerate(vertical_path):
        to_job = path.get("to_job", "")
        years = path.get("years", "")
        nodes.append({"name": to_job, "level": i + 1, "type": "promotion", "years": years})
        edges.append({
            "source": job_name if i == 0 else vertical_path[i-1].get("to_job", ""),
            "target": to_job,
            "years": years
        })
    
    G = nx.DiGraph()
    for node in nodes:
        G.add_node(node["name"], level=node["level"], type=node["type"])
    for edge in edges:
        G.add_edge(edge["source"], edge["target"])
    
    pos = nx.spring_layout(G, seed=42)
    
    fig = go.Figure()
    for edge in edges:
        x0, y0 = pos[edge["source"]]
        x1, y1 = pos[edge["target"]]
        fig.add_trace(go.Scatter(
            x=[x0, x1, None], y=[y0, y1, None],
            mode='lines', line=dict(width=2, color='lightgray'),
            hoverinfo='none', showlegend=False
        ))
    
    for node in nodes:
        x, y = pos[node["name"]]
        color = 'lightgreen' if node["type"] == "current" else 'lightblue'
        size = 40 if node["type"] == "current" else 30
        hover_text = f"<b>{node['name']}</b><br>"
        if node["type"] == "promotion":
            hover_text += f"晋升年限: {node.get('years', 'N/A')}"
        fig.add_trace(go.Scatter(
            x=[x], y=[y], mode='markers+text',
            marker=dict(size=size, color=color, line=dict(width=2, color='darkgray')),
            text=[node["name"]], textposition="top center", textfont=dict(size=10),
            hovertext=hover_text, hoverinfo='text', name=node["name"], showlegend=False
        ))
    
    fig.update_layout(
        title=f"{job_name} - 垂直晋升路径",
        showlegend=False, hovermode='closest',
        margin=dict(b=20, l=5, r=5, t=40),
        annotations=[dict(text="绿色节点: 当前岗位<br>蓝色节点: 晋升岗位<br>连线: 晋升方向",
                         showarrow=False, xref="paper", yref="paper",
                         x=0.005, y=-0.002, xanchor="left", yanchor="bottom", font=dict(size=10))],
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        plot_bgcolor='white'
    )
    return fig

def create_horizontal_change_graph(job_name: str, relation_graph: Dict) -> go.Figure:
    """创建换岗路径图"""
    horizontal_graph = relation_graph.get("horizontal_change_graph", [])
    change_paths = None
    for job in horizontal_graph:
        if job.get("source_job") == job_name:
            change_paths = job.get("change_paths", [])
            break
    
    if not change_paths:
        fig = go.Figure()
        fig.add_annotation(text=f"未找到「{job_name}」的换岗路径数据",
                          xref="paper", yref="paper", x=0.5, y=0.5,
                          showarrow=False, font=dict(size=16, color='gray'))
        return fig
    
    G = nx.Graph()
    G.add_node(job_name, type="center", size=40)
    for path in change_paths:
        target_job = path.get("target_job", "")
        overlap_rate = path.get("skill_overlap_rate", 0)
        difficulty = path.get("change_difficulty", "中等")
        G.add_node(target_job, type="target", size=25, overlap_rate=overlap_rate, difficulty=difficulty)
        G.add_edge(job_name, target_job, overlap_rate=overlap_rate, difficulty=difficulty)
    
    pos = nx.spring_layout(G, k=2, iterations=50, seed=42)
    fig = go.Figure()
    for edge in G.edges(data=True):
        x0, y0 = pos[edge[0]]
        x1, y1 = pos[edge[1]]
        overlap_rate = edge[2].get("overlap_rate", 0)
        width = max(1, overlap_rate * 5)
        fig.add_trace(go.Scatter(
            x=[x0, x1, None], y=[y0, y1, None],
            mode='lines', line=dict(width=width, color='lightblue'),
            hoverinfo='none', showlegend=False
        ))
    
    for node in G.nodes(data=True):
        x, y = pos[node[0]]
        node_type = node[1].get("type", "")
        color = 'lightcoral' if node_type == "center" else 'lightyellow'
        size = node[1].get("size", 30)
        hover_text = f"<b>{node[0]}</b><br>"
        if node_type == "target":
            overlap_rate = node[1].get("overlap_rate", 0)
            difficulty = node[1].get("difficulty", "中等")
            hover_text += f"技能重叠率: {overlap_rate:.0%}<br>换岗难度: {difficulty}"
        fig.add_trace(go.Scatter(
            x=[x], y=[y], mode='markers+text',
            marker=dict(size=size, color=color, line=dict(width=2, color='darkgray')),
            text=[node[0]], textposition="top center", textfont=dict(size=10),
            hovertext=hover_text, hoverinfo='text', name=node[0], showlegend=False
        ))
    
    fig.update_layout(
        title=f"{job_name} - 换岗路径图谱",
        showlegend=False, hovermode='closest',
        margin=dict(b=20, l=5, r=5, t=40),
        annotations=[dict(text="红色节点: 当前岗位<br>黄色节点: 可换岗岗位<br>连线粗细: 技能重叠率",
                         showarrow=False, xref="paper", yref="paper",
                         x=0.005, y=-0.002, xanchor="left", yanchor="bottom", font=dict(size=10))],
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        plot_bgcolor='white'
    )
    return fig

# ========== 辅助函数 ==========
def extract_info_from_messages(keyword: str) -> str:
    """从对话历史中提取包含特定关键词的内容"""
    for msg in reversed(st.session_state.messages):
        if msg["role"] == "assistant" and keyword in msg["content"]:
            return msg["content"]
    return None

def get_last_assistant_message() -> str:
    """获取最后一条助手消息"""
    for msg in reversed(st.session_state.messages):
        if msg["role"] == "assistant":
            return msg["content"]
    return None

def split_assistant_message(message: str) -> Dict[str, str]:
    """
    将智能体回答按部分分割（使用明确的第一/二/三部分标识）

    假设格式：
    ### 第一部分 学生能力画像
    [内容]

    ### 第二部分 岗位匹配分析
    [内容]

    ### 第三部分 职业发展路径
    [内容]

    ### 第四部分 总结与建议
    [内容]
    """
    parts = {
        "student_profile": "",
        "job_match": "",
        "career_path": "",
        "summary": ""
    }

    # 定义明确的分隔符
    part_markers = {
        "student_profile": "### 第一部分",
        "job_match": "### 第二部分",
        "career_path": "### 第三部分",
        "summary": "### 第四部分"
    }

    # 找到各部分的起始位置
    part_positions = {}
    for part_name, marker in part_markers.items():
        pos = message.find(marker)
        if pos != -1:
            part_positions[part_name] = pos

    # 按位置排序
    sorted_parts = sorted(part_positions.items(), key=lambda x: x[1])

    # 提取各部分内容
    for i, (part_name, start_pos) in enumerate(sorted_parts):
        # 查找结束位置（下一个部分的开始）
        if i < len(sorted_parts) - 1:
            end_pos = sorted_parts[i + 1][1]
        else:
            end_pos = len(message)

        # 提取内容（跳过分隔符行）
        content_start = start_pos + len(part_markers[part_name])
        content = message[content_start:end_pos].strip()

        # 移除标题行（如果有）
        lines = content.split('\n')
        if lines and any(marker in lines[0] for markers in [
            "学生能力画像", "岗位匹配分析", "职业发展路径", "总结与建议"
        ] for marker in markers):
            lines = lines[1:]

        parts[part_name] = '\n'.join(lines).strip()

    # 如果没有找到明确分隔符，尝试兼容旧格式
    if not any(parts.values()):
        # 兼容旧格式：使用### 作为分隔符
        section_markers = {
            "student_profile": ["### 学生能力画像", "## 学生能力画像"],
            "job_match": ["### 岗位匹配分析", "## 岗位匹配分析"],
            "career_path": ["### 职业发展路径", "## 职业发展路径"],
            "summary": ["### 总结与建议", "## 总结与建议"]
        }

        # 按行查找章节
        current_section = ""
        current_type = None
        section_order = []

        lines = message.split('\n')
        for line in lines:
            line_stripped = line.strip()
            matched = False

            for section_type, markers in section_markers.items():
                for marker in markers:
                    if line_stripped.startswith(marker):
                        # 保存当前章节
                        if current_type and current_section:
                            parts[current_type] = current_section.strip()

                        # 开始新章节
                        current_section = ""
                        current_type = section_type
                        if section_type not in section_order:
                            section_order.append(section_type)
                        matched = True
                        break

                if matched:
                    break

            if not matched and current_type:
                current_section += line + "\n"

        # 保存最后一个章节
        if current_type and current_section:
            parts[current_type] = current_section.strip()

    return parts

def extract_student_profile(message: str) -> str:
    """从消息中提取学生画像部分"""
    parts = split_assistant_message(message)
    return parts["student_profile"] or message

def extract_job_match(message: str) -> str:
    """从消息中提取岗位匹配部分"""
    parts = split_assistant_message(message)
    return parts["job_match"]

def extract_career_path(message: str) -> str:
    """从消息中提取职业路径部分"""
    parts = split_assistant_message(message)
    return parts["career_path"]

def parse_ability_scores(text: str) -> Dict:
    """从文本中解析能力分数（支持0-10分制）"""
    scores = {
        '专业技能': 0,
        '实习项目能力': 0,
        '证书资质': 0,
        '创新能力': 0,
        '学习能力': 0,
        '抗压能力': 0,
        '沟通协作能力': 0
    }

    ability_mapping = {
        '专业技能能力': '专业技能',
        '专业技能': '专业技能',
        '实习与项目能力': '实习项目能力',
        '实习项目能力': '实习项目能力',
        '证书与资质能力': '证书资质',
        '证书资质能力': '证书资质',
        '证书资质': '证书资质',
        '创新能力': '创新能力',
        '学习能力': '学习能力',
        '抗压能力': '抗压能力',
        '沟通与协作能力': '沟通协作能力',
        '沟通协作能力': '沟通协作能力',
    }

    # 方法1: 解析Markdown表格格式
    lines = text.split('\n')
    for line in lines:
        if any(keyword in line for keyword in ['能力维度', '得分', '点评', '---', '公司名称', '|']):
            continue
        if not line.strip():
            continue

        # 尝试表格格式
        parts = line.split('|')
        parts = [p.strip() for p in parts if p.strip()]

        if len(parts) >= 2:
            ability_name = parts[0].strip()
            if len(parts) >= 2:
                score_part = parts[1].strip()
                match = re.search(r'(\d+\.?\d*)', score_part)
                if match:
                    raw_score = float(match.group(1))
                    # 确保分数在0-10之间
                    final_score = min(10.0, max(0.0, raw_score))
                    mapped_name = ability_mapping.get(ability_name)
                    if mapped_name and mapped_name in scores:
                        scores[mapped_name] = final_score

    # 方法2: 如果表格解析失败，尝试文本格式（例如：专业技能: 7.5分）
    if not any(scores.values()):
        # 匹配 "能力名: 分数分" 或 "能力名: 分数/10" 等格式
        patterns = [
            r'([^\d:：]+)[:：]\s*(\d+\.?\d*)\s*[分分]',  # 专业技能: 7.5分
            r'([^\d/]+?)[:：]\s*(\d+\.?\d*)\s*/\s*10',  # 专业技能: 7.5/10
            r'([^\d.]+?)[:：]\s*(\d+\.?\d*)',  # 专业技能: 7.5
            r'([^\d.]+?)\s*[（(]\s*(\d+\.?\d*)\s*[）)]',  # 专业技能（7.5）
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                ability_name = match[0].strip()
                score = float(match[1])

                # 确保分数在0-10之间
                score = min(10.0, max(0.0, score))

                # 映射能力名称
                mapped_name = None
                for key, value in ability_mapping.items():
                    if key in ability_name or ability_name in key:
                        mapped_name = value
                        break

                if mapped_name and mapped_name in scores:
                    scores[mapped_name] = score

        # 方法3: 尝试提取分数关键词（例如：7.5分）
        if not any(scores.values()):
            # 查找所有数字，判断是否是分数
            all_numbers = re.findall(r'(\d+\.?\d*)\s*分', text)
            if all_numbers:
                # 按能力名称顺序分配分数
                ability_keys = list(ability_mapping.values())
                for i, num in enumerate(all_numbers):
                    if i < len(ability_keys):
                        score = float(num)
                        # 确保分数在0-10之间
                        score = min(10.0, max(0.0, score))
                        scores[ability_keys[i]] = score

    return scores

def read_word_document(file) -> str:
    """读取Word文档内容"""
    try:
        doc = Document(io.BytesIO(file.read()))
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    content.append(row_text)
        return "\n".join(content)
    except Exception as e:
        st.error(f"❌ 读取Word文档失败: {e}")
        return ""

def render_ability_radar_chart_from_text(text: str):
    """从文本渲染学生能力雷达图（0-10分制）"""
    scores = parse_ability_scores(text)

    abilities = [
        {'name': '专业技能', 'value': scores.get('专业技能', 0)},
        {'name': '实习项目', 'value': scores.get('实习项目能力', 0)},
        {'name': '证书资质', 'value': scores.get('证书资质', 0)},
        {'name': '创新能力', 'value': scores.get('创新能力', 0)},
        {'name': '学习能力', 'value': scores.get('学习能力', 0)},
        {'name': '抗压能力', 'value': scores.get('抗压能力', 0)},
        {'name': '沟通协作', 'value': scores.get('沟通协作能力', 0)},
    ]

    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=[a['value'] for a in abilities],
        theta=[a['name'] for a in abilities],
        fill='toself',
        name='能力画像',
        line_color='#667eea',
        fillcolor='rgba(102, 126, 234, 0.3)'
    ))
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 10])),
        showlegend=True,
        title='📊 学生7大能力维度画像（满分10分）',
        height=450,
        margin=dict(l=20, r=20, t=40, b=20)
    )
    st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

    cols = st.columns(4)
    for i, (name, score) in enumerate(scores.items()):
        with cols[i % 4]:
            delta_color = "normal" if score >= 6.0 else "inverse"
            st.metric(
                name,
                f"{score:.1f}/10",
                delta=f"{'优秀' if score >= 8.0 else '良好' if score >= 6.0 else '需提升'}",
                delta_color=delta_color
            )

def call_coze_api(prompt: str, session_id: str, progress_callback=None) -> str:
    """调用Coze智能体API

    Args:
        prompt: 用户提示词
        session_id: 会话ID
        progress_callback: 进度回调函数，用于流式更新响应内容

    Returns:
        完整的响应文本
    """
    headers = {
        "Authorization": f"Bearer {API_TOKEN}",
        "Content-Type": "application/json",
        "Accept": "text/event-stream"
    }

    payload = {
        "content": {
            "query": {
                "prompt": [
                    {
                        "type": "text",
                        "content": {
                            "text": prompt
                        }
                    }
                ]
            }
        },
        "type": "query",
        "session_id": session_id,
        "project_id": PROJECT_ID
    }

    full_response = ""
    try:
        response = requests.post(
            API_URL,
            json=payload,
            headers=headers,
            stream=True,
            timeout=120
        )

        # 添加状态码错误处理
        try:
            response.raise_for_status()
        except Exception as e:
            full_response = f"❌ API调用失败: HTTP {response.status_code}\n\n"
            full_response += f"错误详情: {str(e)}\n\n"
            full_response += f"响应内容:\n{response.text[:500]}"
            return full_response

        # 处理流式响应
        if response.status_code == 200:
            for line in response.iter_lines(decode_unicode=True):
                if line and line.startswith("data:"):
                    data_text = line[5:].strip()
                    try:
                        parsed = json.loads(data_text)

                        # 提取内容
                        content = ""
                        if isinstance(parsed, dict):
                            # 类型1: answer类型
                            if parsed.get('type') == 'answer':
                                content = parsed.get('content', {}).get('answer', '')
                            # 类型2: message字段
                            elif 'message' in parsed:
                                content = parsed['message']
                            # 类型3: content字段（字符串）
                            elif 'content' in parsed and isinstance(parsed['content'], str):
                                content = parsed['content']
                            # 类型4: content字段（字典）
                            elif 'content' in parsed and isinstance(parsed['content'], dict):
                                content = parsed['content'].get('answer', '')
                            # 类型5: text字段
                            elif 'text' in parsed:
                                content = parsed['text']

                            if content:
                                full_response += content

                                # 如果有进度回调函数，调用它来更新显示
                                if progress_callback:
                                    progress_callback(full_response)

                            # 检查是否结束
                            if parsed.get('type') == 'message_end':
                                break
                    except json.JSONDecodeError:
                        # JSON解析失败，直接添加原始文本
                        full_response += data_text
                        if progress_callback:
                            progress_callback(full_response)
        else:
            full_response = f"❌ API调用失败: HTTP {response.status_code}\n\n{response.text}"

    except requests.exceptions.Timeout:
        full_response = "❌ 请求超时，请稍后重试或检查网络连接。"
    except requests.exceptions.ConnectionError:
        full_response = "❌ 网络连接失败，请检查网络设置。"
    except Exception as e:
        full_response = f"❌ 发生错误: {str(e)}"

    return full_response

# ========== 主函数 ==========
def main():
    with st.sidebar:
        # 显示Logo
        try:
            st.image("logo.png", width=250, output_format="PNG")
        except Exception:
            # 如果logo.png不存在，使用emoji作为fallback
            st.markdown("# 🎯 模方职航")

        st.markdown("### 大学生职业规划助手")
        st.markdown("---")
        
        page = st.radio(
            "选择页面",
            ["💬 对话助手", "👤 学生画像", "📈 岗位匹配", "🛤️ 职业路径", "📋 岗位画像"],
            label_visibility="collapsed",
            index=["💬 对话助手", "👤 学生画像", "📈 岗位匹配", "🛤️ 职业路径", "📋 岗位画像"].index(st.session_state.page)
        )

        # 更新session_state中的页面
        if page != st.session_state.page:
            st.session_state.page = page
        
        st.markdown("---")
        
        st.markdown("### 📄 上传简历")
        uploaded_file = st.file_uploader("上传Word简历", type=['docx'], key="resume_uploader")

        if uploaded_file:
            st.success(f"✅ 已上传: {uploaded_file.name}")
            if st.button("📖 读取并分析", key="read_resume"):
                with st.spinner("正在读取简历...（请稍等一到两分钟，不要切换画面）"):
                    resume_content = read_word_document(uploaded_file)
                    if resume_content:
                        st.session_state.resume_content = resume_content

                        # 构建分析提示
                        prompt = f"请分析这份简历，包括学生画像、岗位匹配和职业路径规划：\n\n{resume_content}"

                        # 添加用户消息到对话历史
                        st.session_state.messages.append({"role": "user", "content": prompt})

                        # 调用API获取智能体响应
                        with st.spinner("🤖 智能体正在分析..."):
                            response = call_coze_api(prompt, st.session_state.session_id)

                        # 添加助手响应到对话历史
                        st.session_state.messages.append({"role": "assistant", "content": response})

                        # 跳转到对话框页面
                        st.session_state.page = "💬 对话助手"
                        st.success("✅ 分析完成！请查看对话框中的结果")
                        st.rerun()
                    else:
                        st.error("❌ 简历读取失败，请检查文件格式")
        
        st.markdown("---")
        st.markdown("### ⚙️ 系统状态")
        
        # API连接测试按钮
        if st.button("🔌 测试API连接", key="test_api"):
            with st.spinner("正在测试API连接..."):
                try:
                    test_headers = {
                        "Authorization": f"Bearer {API_TOKEN}",
                        "Content-Type": "application/json"
                    }
                    test_payload = {
                        "content": {
                            "query": {
                                "prompt": [
                                    {
                                        "type": "text",
                                        "content": {"text": "test"}
                                    }
                                ]
                            }
                        },
                        "type": "query",
                        "session_id": st.session_state.session_id,
                        "project_id": PROJECT_ID
                    }
                    test_response = requests.post(
                        API_URL,
                        json=test_payload,
                        headers=test_headers,
                        timeout=10
                    )
                    
                    if test_response.status_code == 200:
                        st.success("✅ API连接正常！")
                    else:
                        st.error(f"❌ API连接失败: HTTP {test_response.status_code}")
                        if test_response.status_code == 401:
                            st.warning("⚠️ 可能是Token过期或无效，请检查API_TOKEN配置")
                        st.code(test_response.text[:200])
                except Exception as e:
                    st.error(f"❌ 连接测试失败: {str(e)}")
        
        st.markdown(f"📝 **会话ID**: `{st.session_state.session_id[:20]}...`")
        st.markdown(f"💬 **对话轮数**: {len(st.session_state.messages) // 2}")
        st.markdown(f"🌐 **API URL**: `{API_URL}`")
        st.markdown(f"🆔 **Project ID**: `{PROJECT_ID}`")
        
        # 显示Token过期提示
        with st.expander("🔧 API配置说明", expanded=False):
            st.markdown("""
            **当前配置：**
            - API_URL: 智能体流式接口地址
            - API_TOKEN: 访问令牌（Bearer认证）
            - PROJECT_ID: 项目ID
            
            **如需更新配置：**
            1. 确保新的智能体已部署完成
            2. 从Coze平台获取最新的API_TOKEN
            3. 更新代码顶部的API_TOKEN变量
            4. 重新运行应用
            
            **常见问题：**
            - 401错误: Token过期或无效 → 更新API_TOKEN
            - 404错误: Project ID错误 → 检查PROJECT_ID
            - 超时错误: 网络问题或服务不可用 → 检查网络连接
            """)
    
    if page == "💬 对话助手":
        st.title("💬 智能对话助手")
        st.markdown("---")
        
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        if prompt := st.chat_input("请输入你的问题或简历内容...", key="chat_input"):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                message_placeholder = st.empty()

                # 检测是否是简历分析，显示友好提示
                is_resume_analysis = any(
                    keyword in prompt.lower()
                    for keyword in ['简历', '分析', '能力', '画像', '评分', '匹配', '职业', '路径', '规划']
                )

                if is_resume_analysis:
                    message_placeholder.markdown("""
🤖 **智能体正在分析中...**

📊 智能体正在查询知识库并分析您的信息，不要切换界面哦可能会发生报错。请稍等1-2分钟...

---
                    """)
                    # 强制刷新UI，让用户看到提示
                    import time
                    time.sleep(0.5)

                # 定义进度回调函数，用于流式更新
                def update_progress(content):
                    message_placeholder.markdown(content)

                # 调用API，传入进度回调
                full_response = call_coze_api(prompt, st.session_state.session_id, update_progress)

            st.session_state.messages.append({"role": "assistant", "content": full_response})
    
    elif page == "👤 学生画像":
        st.title("👤 学生能力画像")
        st.markdown("---")

        # 获取最后一条助手消息
        last_msg = get_last_assistant_message()

        if not last_msg:
            st.info("👋 请先在「对话助手」页面进行对话，系统会自动生成学生画像。")
        else:
            # 提取学生画像部分
            profile_text = extract_student_profile(last_msg)

            if not profile_text:
                st.info("👋 请在「对话助手」页面询问学生画像相关内容，例如：'请分析我的能力'")
                return

            # 显示解析结果
            with st.expander("🔍 解析结果（点击查看）", expanded=True):
                scores = parse_ability_scores(profile_text)
                st.write("解析出的分数（已限制在0-10范围内）：")
                for key, value in scores.items():
                    st.write(f"- {key}: {value:.1f}")

            # 渲染雷达图
            render_ability_radar_chart_from_text(profile_text)

            # 显示原始文本（如果雷达图无法显示）
            if not any(scores.values()):
                st.warning("⚠️ 未能从文本中提取能力分数，显示原始内容：")
                st.markdown(profile_text)
    
    elif page == "📈 岗位匹配":
        st.title("📈 岗位匹配分析")
        st.markdown("---")

        # 获取最后一条助手消息
        last_msg = get_last_assistant_message()

        if not last_msg:
            st.info("👋 请先在「对话助手」页面进行对话，系统会自动生成岗位匹配分析。")
        else:
            # 方法1：查找"### 第二部分"到"### 第三部分"之间的内容
            part2_start = last_msg.find("### 第二部分")
            part3_start = last_msg.find("### 第三部分")

            match_text = ""

            if part2_start != -1 and part3_start != -1 and part2_start < part3_start:
                # 有明确分隔符，提取第二部分到第三部分之间的内容
                match_text = last_msg[part2_start:part3_start].strip()
                # 移除标题行
                lines = match_text.split('\n')
                if lines and "### 第二部分" in lines[0]:
                    lines = lines[1:]
                match_text = '\n'.join(lines).strip()

            if not match_text:
                # 方法2：使用分割函数
                parts = split_assistant_message(last_msg)
                match_text = parts.get("job_match", "")

            if not match_text:
                # 方法3：提取包含"岗位"或"匹配"的段落（排除职业路径部分）
                paragraphs = last_msg.split('\n\n')
                match_related = []
                in_career_section = False

                for para in paragraphs:
                    # 检查是否进入职业路径部分
                    if any(keyword in para for keyword in ['### 第三部分', '### 职业发展路径', '### 职业路径']):
                        in_career_section = True
                        continue

                    if in_career_section:
                        continue

                    # 提取岗位匹配相关内容
                    if any(keyword in para for keyword in ['岗位', '匹配', '推荐', '适合']):
                        match_related.append(para)

                if match_related:
                    match_text = '\n\n'.join(match_related)

            if not match_text:
                # 方法4：最后降级 - 提取包含"岗位"和"匹配"的行
                lines = last_msg.split('\n')
                match_lines = [line for line in lines if '岗位' in line or '匹配' in line]
                if match_lines:
                    match_text = '\n'.join(match_lines[:20])  # 限制前20行

            if not match_text or len(match_text.strip()) < 10:
                st.info("👋 请在「对话助手」页面询问岗位匹配相关内容，例如：'请分析我和Java开发工程师的匹配度'")
                st.info(f"💡 调试信息：消息总长度={len(last_msg)}, 包含'### 第二部分'={'是' if '### 第二部分' in last_msg else '否'}")
            else:
                # 🔧 过滤职业目标相关内容，只保留到"匹配结论"为止
                # 方法1：查找"💡 匹配结论"并提取到该段落结束
                match_conclusion_pattern = "💡 匹配结论"
                conclusion_pos = match_text.find(match_conclusion_pattern)

                if conclusion_pos != -1:
                    # 找到匹配结论，提取到该段落完整结束
                    paragraphs = match_text.split('\n\n')
                    filtered_paragraphs = []

                    for para in paragraphs:
                        filtered_paragraphs.append(para)
                        # 如果当前段落包含匹配结论，则停止添加后续段落
                        if match_conclusion_pattern in para:
                            break

                    match_text = '\n\n'.join(filtered_paragraphs)
                else:
                    # 方法2：如果没有找到匹配结论，使用关键词过滤
                    # 过滤掉包含职业目标关键词的段落
                    paragraphs = match_text.split('\n\n')
                    filtered_paragraphs = []

                    for para in paragraphs:
                        # 跳过职业目标相关段落，遇到第一个就停止
                        if any(keyword in para for keyword in [
                            '结合你的能力画像和匹配情况',
                            '结合你的能力',
                            '职业目标',
                            '短期目标',
                            '中期目标',
                            '长期目标',
                            '### 职业发展路径',
                            '🚀 三、职业发展路径',
                            '基于4263个真实岗位数据',
                            '市场需求量',
                            '社会需求分析'
                        ]):
                            break
                        filtered_paragraphs.append(para)

                    match_text = '\n\n'.join(filtered_paragraphs)

                st.markdown("### 🎯 匹配结果")
                st.markdown(match_text)
    
    elif page == "🛤️ 职业路径":
        st.title("🛤️ 职业发展路径")
        st.markdown("---")

        # 获取最后一条助手消息
        last_msg = get_last_assistant_message()

        if not last_msg:
            st.info("👋 请先在「对话助手」页面进行对话，系统会自动生成职业路径规划。")
        else:
            # 方法1：查找"### 第三部分"之后的所有内容
            part3_start = last_msg.find("### 第三部分")

            path_text = ""

            if part3_start != -1:
                # 有明确分隔符，提取第三部分之后的所有内容
                path_text = last_msg[part3_start:].strip()
                # 移除标题行
                lines = path_text.split('\n')
                if lines and "### 第三部分" in lines[0]:
                    lines = lines[1:]
                path_text = '\n'.join(lines).strip()

                # 移除最后一句"需要我把这份报告导出为可编辑的Word文档吗？"
                if "需要我把这份报告导出为可编辑的Word文档吗？" in path_text:
                    path_text = path_text.replace("需要我把这份报告导出为可编辑的Word文档吗？", "").strip()

            if not path_text:
                # 方法2：使用分割函数
                parts = split_assistant_message(last_msg)
                path_text = parts.get("career_path", "")
                summary = parts.get("summary", "")
                if summary:
                    path_text += "\n\n" + summary

                # 移除最后一句
                if "需要我把这份报告导出为可编辑的Word文档吗？" in path_text:
                    path_text = path_text.replace("需要我把这份报告导出为可编辑的Word文档吗？", "").strip()

            if not path_text:
                # 方法3：提取从职业目标开始的所有内容
                paragraphs = last_msg.split('\n\n')
                path_related = []
                started = False

                for para in paragraphs:
                    # 检测职业目标相关的起始标记
                    if not started:
                        if any(keyword in para for keyword in ['### 第三部分', '### 职业发展路径', '### 职业路径', '### 发展路径', '结合你的能力', '职业目标', '短期目标']):
                            started = True
                            path_related.append(para)
                    else:
                        # 已经开始，继续添加后续段落
                        # 排除不需要的段落
                        if "需要我把这份报告导出为可编辑的Word文档吗？" in para:
                            continue
                        path_related.append(para)

                if path_related:
                    path_text = '\n\n'.join(path_related)

                    # 移除最后一句
                    if "需要我把这份报告导出为可编辑的Word文档吗？" in path_text:
                        path_text = path_text.replace("需要我把这份报告导出为可编辑的Word文档吗？", "").strip()

            if not path_text:
                # 方法4：最后降级 - 提取包含路径、建议、提升等关键词的行
                lines = last_msg.split('\n')
                path_lines = [line for line in lines if any(k in line for k in ['路径', '建议', '提升', '规划', '发展', '晋升', '目标'])]
                if path_lines:
                    path_text = '\n'.join(path_lines[:30])  # 限制前30行

                    # 移除最后一句
                    if "需要我把这份报告导出为可编辑的Word文档吗？" in path_text:
                        path_text = path_text.replace("需要我把这份报告导出为可编辑的Word文档吗？", "").strip()

            if not path_text or len(path_text.strip()) < 10:
                st.info("👋 请在「对话助手」页面询问职业发展路径相关内容，例如：'请帮我规划职业发展路径'")
                st.info(f"💡 调试信息：消息总长度={len(last_msg)}, 包含'### 第三部分'={'是' if '### 第三部分' in last_msg else '否'}")
            else:
                st.markdown("### 🛤️ 发展路径")
                st.markdown(path_text)
    
    elif page == "📋 岗位画像":
        st.title("📋 岗位职业路径可视化系统")
        st.markdown("---")
        
        # 加载数据
        with st.spinner("正在加载数据..."):
            job_portraits, job_relation_graph, vertical_paths = load_job_data()
        
        if not job_portraits:
            st.error("❌ 未找到岗位数据文件，请确保 assets/ 目录下有以下文件：")
            st.warning("- job_portraits.json")
            st.warning("- job_relation_graph.json")
            st.warning("- job_vertical_paths.json")
            return
        
        st.info("📊 本系统提供岗位画像、垂直晋升路径、换岗路径的可视化展示，帮助您清晰了解岗位发展前景。")
        
        # 岗位选择（右上角）
        all_jobs = get_all_jobs(job_portraits)
        
        # 创建右上角布局
        col_search, col_select = st.columns([1, 2])
        
        with col_search:
            search_query = st.text_input("🔍 搜索岗位", placeholder="输入岗位名称...", key="job_search")
        
        with col_select:
            # 根据搜索过滤岗位
            if search_query:
                filtered_jobs = [job for job in all_jobs if search_query.lower() in job.lower()]
            else:
                filtered_jobs = all_jobs
            
            selected_job = st.selectbox(
                "📋 选择岗位",
                options=filtered_jobs,
                index=0 if filtered_jobs else None,
                key="job_select",
                label_visibility="visible"
            )
        
        if not selected_job:
            st.warning("⚠️ 请先选择一个岗位进行分析")
            return
        
        # 标签页
        tab1, tab2, tab3, tab4 = st.tabs([
            "📋 岗位画像",
            "📈 垂直晋升路径",
            "🔄 换岗路径",
            "🔗 关联图谱"
        ])
        
        # 查找岗位画像
        job_portrait = None
        for job in job_portraits:
            if job.get("standard_job_name") == selected_job:
                job_portrait = job
                break
        
        # 标签页1：岗位画像
        with tab1:
            if job_portrait:
                st.subheader("📋 岗位画像详情")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("岗位名称", job_portrait.get("standard_job_name", ""))
                with col2:
                    st.metric("职级", job_portrait.get("job_level", "未定义"))
                with col3:
                    st.metric("行业", job_portrait.get("industry", "未定义"))
                
                st.write("**学历和专业要求：**")
                education = job_portrait.get("education_required", [])
                major = job_portrait.get("major_required", [])
                col1, col2 = st.columns(2)
                with col1:
                    st.write("学历要求：", ", ".join(education) if education else "未定义")
                with col2:
                    st.write("专业要求：", ", ".join(major) if major else "未定义")
                
                st.write("**核心技能：**")
                skills = job_portrait.get("skills", {})
                common_skills = skills.get("common_skills", [])
                advanced_skills = skills.get("advanced_skills", [])
                
                if common_skills:
                    st.write("**通用技能：**")
                    st.write(", ".join(common_skills[:10]))
                
                if advanced_skills:
                    st.write("**进阶技能：**")
                    st.write(", ".join(advanced_skills[:10]))
                
                col1, col2 = st.columns(2)
                with col1:
                    fig_radar = create_job_radar(job_portrait)
                    st.plotly_chart(fig_radar, use_container_width=True, key=f"radar_{selected_job}")
                
                with col2:
                    fig_salary = create_salary_chart(job_portrait)
                    st.plotly_chart(fig_salary, use_container_width=True, key=f"salary_{selected_job}")
            else:
                st.warning(f"未找到「{selected_job}」的详细画像信息")
        
        # 标签页2：垂直晋升路径
        with tab2:
            st.subheader("📈 垂直晋升路径")
            
            fig_vertical = create_vertical_path_graph(selected_job, vertical_paths)
            st.plotly_chart(fig_vertical, use_container_width=True, key="vertical_main")
            
            st.write("**晋升路径详情：**")
            jobs_data = vertical_paths.get("jobs", {})
            current_job = jobs_data.get(selected_job, {})
            career_path = current_job.get("career_path", {})
            vertical_path = career_path.get("vertical_path", [])
            
            if vertical_path:
                for i, path in enumerate(vertical_path, 1):
                    with st.expander(f"{i}. {path.get('to_job', '')}", expanded=True):
                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric("晋升年限", path.get('years', 'N/A'))
                        with col2:
                            level = path.get('to_level', 0)
                            st.metric("职级", f"{level}")
                        
                        st.write("**晋升要求：**")
                        requirements = path.get('requirements', [])
                        if requirements:
                            for req in requirements:
                                st.write(f"• {req}")
                        else:
                            st.write("暂无详细要求")
                        
                        st.write("**需要学习的技能：**")
                        skills = path.get('skills_to_learn', [])
                        if skills:
                            for skill in skills:
                                st.write(f"• {skill}")
                        else:
                            st.write("暂无技能清单")
            else:
                st.warning(f"暂未找到「{selected_job}」的晋升路径数据")
        
        # 标签页3：换岗路径
        with tab3:
            st.subheader("🔄 换岗路径")
            
            fig_horizontal = create_horizontal_change_graph(selected_job, job_relation_graph)
            st.plotly_chart(fig_horizontal, use_container_width=True, key="horizontal_main")
            
            st.write("**可换岗岗位详情：**")
            horizontal_graph = job_relation_graph.get("horizontal_change_graph", [])
            
            change_paths = None
            for job in horizontal_graph:
                if job.get("source_job") == selected_job:
                    change_paths = job.get("change_paths", [])
                    break
            
            if change_paths:
                for i, path in enumerate(change_paths, 1):
                    with st.expander(f"{i}. {path.get('target_job', '')}", expanded=True):
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            overlap_rate = path.get('skill_overlap_rate', 0)
                            st.metric("技能重叠率", f"{overlap_rate:.0%}")
                        with col2:
                            difficulty = path.get('change_difficulty', '中等')
                            st.metric("换岗难度", difficulty)
                        with col3:
                            if overlap_rate >= 0.9:
                                success = "极易"
                            elif overlap_rate >= 0.8:
                                success = "容易"
                            elif overlap_rate >= 0.7:
                                success = "中等"
                            else:
                                success = "较难"
                            st.metric("成功率", success)
                        
                        if difficulty == "低":
                            st.success("换岗难度低，建议优先考虑")
                        elif difficulty == "中":
                            st.info("换岗难度中等，需要一定准备")
                        else:
                            st.warning("换岗难度较高，需要充分准备")
            else:
                st.warning(f"暂未找到「{selected_job}」的换岗路径数据")
        
        # 标签页4：关联图谱
        with tab4:
            st.subheader("🔗 岗位关联图谱")
            st.info("此视图展示当前岗位的所有关联关系，包括晋升路径和换岗路径")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("**垂直晋升方向：**")
                st.plotly_chart(fig_vertical, use_container_width=True, key="vertical_summary")
            
            with col2:
                st.write("**横向换岗方向：**")
                st.plotly_chart(fig_horizontal, use_container_width=True, key="horizontal_summary")
            
            st.write("**发展路径总结：**")
            
            jobs_data = vertical_paths.get("jobs", {})
            current_job = jobs_data.get(selected_job, {})
            career_path = current_job.get("career_path", {})
            vertical_path = career_path.get("vertical_path", [])
            
            horizontal_graph = job_relation_graph.get("horizontal_change_graph", [])
            change_paths = None
            for job in horizontal_graph:
                if job.get("source_job") == selected_job:
                    change_paths = job.get("change_paths", [])
                    break
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("垂直晋升路径", f"{len(vertical_path)} 步")
            with col2:
                st.metric("横向换岗方向", f"{len(change_paths) if change_paths else 0} 个")
            with col3:
                if change_paths:
                    avg_overlap = sum(p.get('skill_overlap_rate', 0) for p in change_paths) / len(change_paths)
                    st.metric("平均技能重叠率", f"{avg_overlap:.0%}")
                else:
                    st.metric("平均技能重叠率", "N/A")

if __name__ == "__main__":
    main()
